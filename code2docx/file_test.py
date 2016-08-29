import os

import math
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import argparse

# 页码
# 去掉注释

parser = argparse.ArgumentParser()
parser.add_argument('-p', dest='project_path', help='the absolutely path of your project')
parser.add_argument('-t', dest='file_type', action='append', help='the file types your project contains')
args = parser.parse_args()

PATH = args.project_path
TYPES = args.file_type

PROJECT_PATH = os.path.abspath(os.curdir)
STATIC_PATH = os.path.join(PROJECT_PATH, 'static')
DOCX_PATH = os.path.join(STATIC_PATH, 'docx')
LINE_LEN = 82

if __name__ == '__main__':
    if PATH is None or not os.path.exists(PATH) or not os.path.isabs(PATH) or os.path.isfile(PATH):
        print("plz input an absolutely path of your project")
        exit()
    if TYPES is None or len(TYPES) == 0:
        print("plz input the file types your project contains")
        exit()

    document = Document()
    styles = document.styles

    style_content = styles.add_style('ContentStyle', WD_STYLE_TYPE.PARAGRAPH)
    style_content.base_style = styles['Normal']
    cont_font = style_content.font
    cont_font.name = '宋体'
    cont_font.size = Pt(10.5)  # 10.5 五号
    para_cont_format = style_content.paragraph_format
    para_cont_format.line_spacing = Pt(18)  # 单倍行距 行间距18

    style_title = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    style_content.base_style = styles['ContentStyle']  # didn't work
    title_font = style_title.font
    title_font.name = '宋体'
    title_font.size = Pt(10.5)
    para_title_format = style_title.paragraph_format
    para_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para_title_format.line_spacing = Pt(18)

    list_dirs = os.walk(PATH)

    for root, dirs, files in list_dirs:
        for f in files:
            if os.path.splitext(f)[1].replace('.', '') in TYPES:
                abs_path = os.path.join(root, f)
                with open(abs_path) as fp:
                    content = fp.read()
                    left_num = math.ceil((LINE_LEN - len(f)) * 1.0 / 2.0)
                    right_num = left_num - 1
                    title_str = left_num * '=' + f + right_num * '='
                    paragraph = document.add_paragraph(title_str, 'TitleStyle')
                    paragraph = document.add_paragraph(content, 'ContentStyle')
                    if f != files[-1]:
                        document.add_page_break()

    document.save(os.path.join(DOCX_PATH, 'demo.docx'))
